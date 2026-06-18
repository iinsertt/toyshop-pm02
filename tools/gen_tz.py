"""Генерация документа «Техническое задание.docx» для проекта ToyShop (Этап 1).

Документ оформляется шрифтом Times New Roman 14 пт с полуторным интервалом и
содержит пять обязательных разделов, включая встроенные UML-диаграммы.

Запуск: ``python tools/gen_tz.py``
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(ROOT, "docs")

FONT_NAME = "Times New Roman"
FONT_SIZE = 14


def setup_styles(document):
    """Настроить базовый стиль документа по требованиям (Times New Roman 14, 1.5)."""
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)
    # Для корректного применения шрифта к кириллице.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)

    for section in document.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def add_heading(document, text, level=1):
    """Добавить заголовок раздела (полужирный, Times New Roman)."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(16 if level == 1 else 14)
    return paragraph


def add_body(document, text):
    """Добавить абзац основного текста с красной строкой и выравниванием по ширине."""
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    return paragraph


def add_bullet(document, text):
    """Добавить пункт маркированного списка."""
    paragraph = document.add_paragraph(text, style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_image(document, filename, width_cm=16):
    """Вставить изображение по центру."""
    path = os.path.join(DOCS_DIR, filename)
    document.add_picture(path, width=Cm(width_cm))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_title_page(document):
    """Сформировать титульный лист технического задания."""
    for _ in range(6):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ТЕХНИЧЕСКОЕ ЗАДАНИЕ")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(28)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        "на разработку информационной системы\n«Интернет-магазин игрушек ToyShop»"
    )
    sub_run.font.name = FONT_NAME
    sub_run.font.size = Pt(16)

    for _ in range(10):
        document.add_paragraph()

    for text in (
        "Дисциплина: ПМ.02 «Осуществление интеграции программных модулей»",
        "Вариант 5 — магазин игрушек",
        "Выполнил: Подмарьков Роман Михайлович",
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)

    for _ in range(6):
        document.add_paragraph()
    year = document.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year.add_run("2026").font.size = Pt(14)
    document.add_page_break()


def build_document():
    """Собрать полный документ технического задания."""
    document = Document()
    setup_styles(document)
    build_title_page(document)

    # Раздел 1
    add_heading(document, "1 Описание предметной области")
    add_body(
        document,
        "Информационная система «ToyShop» предназначена для автоматизации работы "
        "интернет-магазина игрушек. Система позволяет покупателю просматривать каталог "
        "товаров, формировать корзину и оформлять заказы, а администратору — управлять "
        "товарами и обрабатывать поступившие заказы.",
    )
    add_body(document, "В магазине представлены товары четырёх категорий:")
    for item in ("конструкторы;", "куклы;", "машинки;", "настольные игры."):
        add_bullet(document, item)
    add_body(document, "В системе предусмотрены две роли пользователей:")
    add_bullet(document, "Покупатель — просматривает каталог, ищет и фильтрует товары, "
                         "управляет корзиной, оформляет заказ;")
    add_bullet(document, "Администратор — управляет каталогом товаров (добавление, "
                         "редактирование, удаление) и заказами (просмотр, изменение статуса).")

    # Раздел 2
    add_heading(document, "2 Функциональные требования")
    add_body(document, "Система должна обеспечивать выполнение следующих функций.")
    add_heading(document, "2.1 Каталог товаров", level=2)
    for item in (
        "загрузка списка товаров из JSON-файла;",
        "отображение всех товаров с указанием цены, остатка и описания;",
        "поиск товаров по названию;",
        "фильтрация товаров по категории и диапазону цены.",
    ):
        add_bullet(document, item)
    add_heading(document, "2.2 Корзина", level=2)
    for item in (
        "добавление товара в корзину;",
        "удаление товара из корзины;",
        "изменение количества товара;",
        "автоматический расчёт общей суммы заказа.",
    ):
        add_bullet(document, item)
    add_heading(document, "2.3 Оформление заказа", level=2)
    for item in (
        "сбор данных покупателя (фамилия, имя, адрес, телефон, e-mail);",
        "валидация введённых данных;",
        "генерация уникального номера заказа формата ORD-ГГГГММДД-NNNN;",
        "сохранение заказа в базу данных SQLite;",
        "вывод подтверждения заказа на экран.",
    ):
        add_bullet(document, item)
    add_heading(document, "2.4 Административная панель", level=2)
    for item in (
        "просмотр списка всех заказов;",
        "изменение статуса заказа (новый, в обработке, доставлен, отменён);",
        "добавление новых товаров;",
        "редактирование и удаление существующих товаров.",
    ):
        add_bullet(document, item)

    # Раздел 3
    add_heading(document, "3 Нефункциональные требования")
    add_heading(document, "3.1 Быстродействие", level=2)
    add_body(document, "Время отклика системы на любую операцию пользователя не должно "
                       "превышать одной секунды при объёме каталога до нескольких тысяч "
                       "товаров.")
    add_heading(document, "3.2 Надёжность и обработка ошибок", level=2)
    add_body(document, "Все операции, способные привести к сбою (чтение файлов, работа с "
                       "базой данных, ввод пользователя), защищены конструкциями try/except. "
                       "Некорректные данные отклоняются на этапе валидации. Все ключевые "
                       "события и ошибки фиксируются в журнале (файл logs/app.log).")
    add_heading(document, "3.3 Масштабируемость", level=2)
    add_body(document, "Модульная архитектура позволяет расширять систему без переработки "
                       "существующего кода: добавлять новые модули, заменять хранилище "
                       "каталога (JSON) на полноценную СУБД, наращивать число категорий и "
                       "товаров.")

    # Раздел 4
    add_heading(document, "4 Архитектура системы")
    add_body(document, "Система построена по модульному принципу и состоит из пяти модулей. "
                       "Модуль интеграции (main.py) предоставляет консольное меню и связывает "
                       "остальные модули. Общая инфраструктура (модели данных, исключения, "
                       "логирование) вынесена в пакет src/common. Каталог товаров хранится в "
                       "файле products.json, заказы — в базе данных SQLite (orders.db).")
    add_body(document, "Потоки данных: каталог передаёт выбранный товар в корзину; корзина "
                       "передаёт позиции и сумму в модуль оформления заказа; оформленный заказ "
                       "сохраняется в базу данных; административная панель читает заказы из базы "
                       "и изменяет каталог товаров.")
    add_body(document, "Диаграмма компонентов системы приведена на рисунке 1.")
    add_image(document, "component_diagram.png")
    caption1 = document.add_paragraph("Рисунок 1 — Диаграмма компонентов системы")
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Раздел 5
    add_heading(document, "5 Диаграмма вариантов использования")
    add_body(document, "В системе выделены два актёра — Покупатель и Администратор. "
                       "Диаграмма вариантов использования приведена на рисунке 2.")
    add_image(document, "use_case_diagram.png")
    caption2 = document.add_paragraph("Рисунок 2 — Диаграмма вариантов использования")
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "5.1 Основные и альтернативные сценарии", level=2)
    add_body(document, "Основной сценарий «Оформление заказа»: покупатель добавляет товары "
                       "в корзину, переходит к оформлению, вводит корректные данные, система "
                       "проверяет их, сохраняет заказ и выводит его номер.")
    add_body(document, "Альтернативный сценарий «Ошибка валидации»: при вводе некорректного "
                       "телефона или e-mail система отклоняет данные и выводит сообщение об "
                       "ошибке, не сохраняя заказ; покупатель повторяет ввод.")
    add_body(document, "Альтернативный сценарий «Пустая корзина»: при попытке оформить заказ "
                       "с пустой корзиной система сообщает, что оформлять нечего.")
    add_body(document, "Основной сценарий «Управление товарами»: администратор добавляет, "
                       "редактирует или удаляет товар; изменения немедленно отражаются в "
                       "каталоге. Альтернативный сценарий: при попытке изменить несуществующий "
                       "товар система сообщает об ошибке.")

    output = os.path.join(DOCS_DIR, "technical_specification.docx")
    document.save(output)
    print(f"Создано: {output}")


if __name__ == "__main__":
    build_document()
